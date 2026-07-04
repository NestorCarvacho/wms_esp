#!/usr/bin/env python3
"""Genera documentación de proyecto Khepri Software (estilo DUOC/PMI)."""
from __future__ import annotations

import shutil
import sys
from datetime import date, timedelta
from pathlib import Path

SCRIPTS = Path(__file__).resolve().parent
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from generar_requerimientos_excel import HEADERS, ROWS, main as gen_requerimientos

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "proyecto"
OUT.mkdir(parents=True, exist_ok=True)

APP = "Khepri Software"
TAGLINE = "Tu WMS a tu medida"
FECHA = date(2026, 5, 31)


def _hdr(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str, bold: bool = False) -> None:
    run = doc.add_paragraph().add_run(text)
    run.bold = bold


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val


def gen_acta() -> Path:
    path = OUT / "01_acta_constitucion_khepri.docx"
    doc = Document()
    title = doc.add_heading("Acta de Constitución del Proyecto", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _hdr(doc, "Información del proyecto", 2)
    _table(
        doc,
        ["Campo", "Valor"],
        [
            ["Empresa / Organización", "Khepri Software"],
            ["Proyecto", f"{APP} — {TAGLINE}"],
            ["Fecha de preparación", FECHA.strftime("%d/%m/%Y")],
            ["Cliente", "Organizaciones con operación logística multi-bodega"],
            ["Patrocinador principal", "Gerencia / Dueño del producto SaaS"],
            ["Gerente de proyecto", "Por definir"],
        ],
    )

    _hdr(doc, "Propósito y justificación", 2)
    _p(
        doc,
        "Desarrollar una plataforma WMS multi-empresa en la nube que permita controlar "
        "stock por ubicación, movimientos auditados (recepción, traslado, despacho) y "
        "administración de usuarios con permisos granulares, reduciendo errores en piso "
        "y mejorando la trazabilidad operativa.",
    )

    _hdr(doc, "Descripción y entregables", 2)
    for item in [
        "API REST FastAPI con autenticación JWT y RBAC.",
        "Frontend React con módulos de catálogo, inventario y administración.",
        "Base de datos MySQL multi-tenant.",
        "Despliegue en Railway (API + BD + frontend).",
        "Documentación ERS, casos de uso, EDT y matriz de trazabilidad.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _hdr(doc, "Objetivos e indicadores", 2)
    _table(
        doc,
        ["Objetivo", "Indicador de éxito"],
        [
            ["Trazabilidad operativa", "100% movimientos con usuario, fecha y zona registrados"],
            ["Multi-empresa", "Al menos 2 tenants operando con aislamiento de datos"],
            ["Adopción en piso", "Operadores registran recepciones/traslados/despachos en sistema"],
            ["Disponibilidad", "API /health operativo en producción"],
        ],
    )

    _hdr(doc, "Cronograma (referencial)", 2)
    _table(
        doc,
        ["Fase", "Inicio", "Término"],
        [
            ["Análisis y diseño", "01/2026", "02/2026"],
            ["Desarrollo MVP (auth, catálogo, inventario)", "02/2026", "04/2026"],
            ["Pruebas, despliegue Railway", "04/2026", "05/2026"],
            ["Cierre y documentación", "05/2026", "06/2026"],
        ],
    )

    _hdr(doc, "Presupuesto inicial (referencial)", 2)
    _p(doc, "A definir según tarifas de desarrollo, hosting Railway y dominio/correo transaccional.")

    _hdr(doc, "Riesgos iniciales", 2)
    for r in [
        "Resistencia al cambio en operadores de bodega.",
        "Configuración incorrecta de permisos RBAC en empresas nuevas.",
        "Dependencia de servicios externos (Resend, Railway).",
        "Ampliación de alcance (órdenes de compra/venta).",
    ]:
        doc.add_paragraph(r, style="List Bullet")

    _hdr(doc, "Stakeholders", 2)
    _table(
        doc,
        ["Nombre / Rol", "Interés"],
        [
            ["Super admin (empresa maestra)", "Administrar tenants y configuración global"],
            ["Administrador de empresa", "Usuarios, catálogo, permisos"],
            ["Operador de bodega", "Recepción, traslado, despacho, consulta stock"],
            ["Equipo de desarrollo", "Implementación y despliegue"],
        ],
    )

    _hdr(doc, "Aprobaciones", 2)
    _table(doc, ["Rol", "Fecha", "Firma"], [["Patrocinador", "", ""], ["Gerente de proyecto", "", ""]])

    doc.save(path)
    return path


def gen_ers() -> Path:
    path = OUT / "02_ers_khepri_software.docx"
    doc = Document()
    doc.add_heading(f"ERS — {APP}", 0)
    _p(doc, f"Revisión 01 — {FECHA.strftime('%d/%m/%Y')}")
    _p(doc, "Especificación de Requisitos de Software (referencia IEEE 830 / PMI)")

    _hdr(doc, "1. Introducción", 1)
    _hdr(doc, "1.1 Propósito", 2)
    _p(doc, f"Definir requisitos funcionales y no funcionales de {APP} para diseño, desarrollo y pruebas.")

    _hdr(doc, "1.2 Ámbito del sistema", 2)
    _p(
        doc,
        f"{APP} es un WMS multi-tenant en la nube. Administra empresas, usuarios, catálogo "
        "(productos, bodegas, zonas), inventario operativo con movimientos auditados, "
        "reportes exportables y control de acceso por roles.",
    )

    _hdr(doc, "1.3 Definiciones", 2)
    _p(doc, "WMS=Warehouse Management System; RBAC=Control de acceso basado en roles; "
           "Tenant=Empresa cliente; JWT=JSON Web Token; EDT=Estructura de descomposición del trabajo.")

    _hdr(doc, "2. Descripción general", 1)
    _hdr(doc, "2.1 Perspectiva del producto", 2)
    _p(doc, "Producto SaaS independiente; empresa maestra administra tenants. Acceso vía navegador.")

    _hdr(doc, "2.2 Funciones del producto", 2)
    for f in [
        "Autenticación, recuperación de contraseña y bloqueo por intentos.",
        "CRUD multi-empresa: usuarios, cargos, roles, permisos.",
        "Catálogo: productos, bodegas, zonas, unidades.",
        "Inventario: stock por zona, recepción, traslado, despacho, historial.",
        "Dashboard, exportación Excel/PDF, carga masiva productos.",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    _hdr(doc, "2.3 Características de usuarios", 2)
    _table(
        doc,
        ["Perfil", "Conocimientos", "Funciones principales"],
        [
            ["Super admin", "Gestión sistemas", "Empresas, RBAC global, selector tenant"],
            ["Administrador", "Básicos PC/navegador", "Usuarios, catálogo, permisos"],
            ["Operador", "Básicos; escaneo SKU", "Movimientos de inventario en piso"],
            ["Visitante", "Navegación web", "Landing e inicio de sesión"],
        ],
    )

    _hdr(doc, "2.4 Restricciones", 2)
    for c in [
        "Stack: Python FastAPI, React, MySQL.",
        "Despliegue objetivo: Railway.",
        "Correo transaccional: Resend (dominio propio en producción).",
        "Navegadores modernos (Chrome, Edge, Firefox).",
    ]:
        doc.add_paragraph(c, style="List Bullet")

    _hdr(doc, "3. Requisitos específicos", 1)
    _hdr(doc, "3.1 Requisitos funcionales", 2)
    func = [r for r in ROWS if r[2] == "Funcional"]
    _table(
        doc,
        ["ID", "Nombre", "Descripción", "Actores"],
        [[r[0], r[1], r[5], r[4]] for r in func],
    )

    _hdr(doc, "3.2 Requisitos no funcionales", 2)
    nfunc = [r for r in ROWS if r[2] == "No Funcional"]
    _table(
        doc,
        ["ID", "Nombre", "Clasificación", "Descripción"],
        [[r[0], r[1], r[3], r[5]] for r in nfunc],
    )

    _hdr(doc, "4. Anexos", 1)
    for a in [
        "Anexo A: Acta de constitución (01_acta_constitucion_khepri.docx)",
        "Anexo B: Planilla de requerimientos (requerimientos_khepri_software.xlsx)",
        "Anexo C: EDT (03_edt_khepri_software.xlsx)",
        "Anexo D: Casos de uso (04_casos_uso_khepri.docx)",
        "Anexo E: Carta Gantt (05_carta_gantt_khepri.xlsx)",
        "Anexo F: Matriz trazabilidad (06_matriz_trazabilidad_khepri.xlsx)",
        "Anexo G: Catálogo de pantallas (07_catalogo_pantallas_khepri.docx)",
    ]:
        doc.add_paragraph(a, style="List Bullet")

    doc.save(path)
    return path


def _cu_extendido(
    doc: Document,
    codigo: str,
    nombre: str,
    actores: str,
    objetivo: str,
    pre: str,
    post: str,
    pasos: list[tuple[str, str]],
    alternativos: list[tuple[str, str]] | None = None,
) -> None:
    _hdr(doc, f"Caso de Uso {codigo} — {nombre}", 3)
    _p(doc, f"Actores: {actores}")
    _p(doc, f"Objetivo: {objetivo}")
    _p(doc, f"Precondición: {pre}")
    _p(doc, f"Postcondición: {post}")
    _p(doc, "Curso normal:", bold=True)
    _table(doc, ["Paso", "Acción del actor", "Respuesta del sistema"], [
        [str(i + 1), a, s] for i, (a, s) in enumerate(pasos)
    ])
    if alternativos:
        _p(doc, "Cursos alternativos:", bold=True)
        _table(doc, ["Condición", "Respuesta del sistema"], alternativos)


def gen_casos_uso() -> Path:
    path = OUT / "04_casos_uso_khepri.docx"
    doc = Document()
    doc.add_heading(f"Casos de Uso — {APP}", 0)

    _hdr(doc, "Casos de uso — Alto nivel", 1)
    casos_alto = [
        ("CU001", "Autenticar usuario", "Administrador, Operador, Super admin",
         "Seguridad, Operatividad", "R.01, R.03, R.49, R.50"),
        ("CU002", "Restablecer contraseña (admin)", "Administrador",
         "Seguridad, Fiabilidad", "R.02, R.28, R.26"),
        ("CU003", "Gestionar empresas", "Super admin",
         "Seguridad, Mantenibilidad", "R.04, R.05, R.06, R.46"),
        ("CU004", "Gestionar usuarios y RBAC", "Administrador, Super admin",
         "Seguridad", "R.07–R.10, R.37, R.38, R.45"),
        ("CU005", "Gestionar catálogo", "Administrador, Super admin",
         "Operatividad", "R.11–R.17, R.34"),
        ("CU006", "Consultar stock", "Operador, Administrador",
         "Rendimiento", "R.18"),
        ("CU007", "Operar inventario (recepción/traslado/despacho)", "Operador",
         "Precisión, Trazabilidad", "R.19, R.29–R.33, R.35"),
        ("CU008", "Consultar stock y movimientos", "Administrador, Operador",
         "Usabilidad, Rendimiento", "R.18, R.20"),
        ("CU009", "Exportar reportes", "Administrador, Operador",
         "Rendimiento", "R.21, R.47"),
        ("CU010", "Gestionar perfil", "Administrador, Operador",
         "Seguridad", "R.22, R.36"),
        ("CU011", "Acceder landing pública", "Visitante", "Usabilidad", "R.23"),
        ("CU012", "Administrar permisos por cargo", "Administrador, Super admin",
         "Seguridad", "R.10, R.38"),
    ]
    _table(
        doc,
        ["CU", "Nombre", "Actores", "Req. no funcionales", "Requerimientos"],
        casos_alto,
    )

    _hdr(doc, "Casos de uso extendidos", 1)
    _cu_extendido(
        doc, "CU001.1", "Iniciar sesión",
        "Administrador, Operador, Super admin",
        "Acceder al sistema con credenciales válidas.",
        "Usuario registrado y activo; empresa operativa.",
        "Sesión iniciada; JWT emitido con permisos.",
        [
            ("Ingresa email y contraseña en /login y envía formulario.",
             "Valida credenciales, estado de cuenta y empresa; genera token JWT."),
            ("—", "Redirige al panel principal /app con menú según permisos."),
        ],
        [
            ("Credenciales incorrectas", "Muestra error genérico; incrementa intentos fallidos."),
            ("Cuenta bloqueada", "Informa bloqueo temporal o permanente."),
            ("Empresa inhabilitada", "Deniega acceso al tenant."),
        ],
    )
    _cu_extendido(
        doc, "CU002.1", "Restablecer contraseña por administrador",
        "Administrador",
        "Asignar contraseña temporal a un usuario del tenant.",
        "Permiso usuarios.editar; usuario destino activo.",
        "Contraseña actualizada; usuario puede iniciar sesión.",
        [
            ("En Usuarios edita el usuario y define nueva contraseña.", "Valida política de contraseña; persiste hash."),
            ("—", "Usuario cambia contraseña en Mi perfil si lo desea."),
        ],
    )
    _cu_extendido(
        doc, "CU007.1", "Recepcionar mercadería",
        "Operador, Administrador",
        "Incrementar stock en zona de recepción.",
        "Usuario autenticado con permiso inventario.recepcionar; bodega y producto existen.",
        "Stock actualizado; movimiento RECEPCION registrado.",
        [
            ("Selecciona bodega, zona y producto (SKU o búsqueda).", "Muestra stock actual y presentaciones."),
            ("Ingresa cantidad y confirma recepción.", "Valida datos; actualiza stock_zona; crea movimiento_inventario."),
            ("—", "Muestra mensaje de éxito en pantalla."),
        ],
    )
    _cu_extendido(
        doc, "CU003.1", "Inhabilitar empresa",
        "Super admin",
        "Suspender operación de un tenant sin borrar datos.",
        "Usuario super admin; empresa no es maestra.",
        "Empresa con esta_activa=false; excluida de listados agregados.",
        [
            ("En Empresas selecciona Inhabilitar y confirma.", "Marca empresa inhabilitada; conserva datos relacionados."),
            ("—", "Empresa visible solo en selector explícito; usuarios del tenant no pueden login."),
        ],
    )
    _cu_extendido(
        doc, "CU005.1", "Registrar producto",
        "Administrador, Super admin",
        "Agregar SKU al catálogo de la empresa.",
        "Permiso productos.crear; empresa destino operativa.",
        "Producto creado y visible en listado.",
        [
            ("Abre panel Nueva producto e ingresa SKU, nombre, unidad, etc.", "Valida unicidad y campos obligatorios."),
            ("Guarda.", "Persiste producto; cierra panel y actualiza tabla."),
        ],
    )

    doc.save(path)
    return path


EDT_ITEMS = [
    ("1.0", "Khepri Software — WMS multi-tenant", "Proyecto completo"),
    ("1.1", "Autenticación y seguridad", "Login, JWT, cambio contraseña, bloqueo"),
    ("1.2", "Multi-empresa y RBAC", "Empresas, usuarios, roles, permisos, cargos"),
    ("1.3", "Catálogo maestro", "Productos, bodegas, zonas, unidades, tipos"),
    ("1.4", "Inventario operativo", "Stock, recepción, traslado, despacho, historial"),
    ("1.5", "Reportes y exportaciones", "Export Excel/PDF stock y movimientos"),
    ("1.6", "Frontend y experiencia de usuario", "Landing, temas, tablas CRUD, regionalización"),
    ("1.7", "Infraestructura y despliegue", "Railway, MySQL, migraciones, CI GitHub"),
    ("1.8", "Pruebas y QA", "Pruebas funcionales y UAT"),
    ("1.9", "Migración a producción", "Puesta en marcha Railway"),
]


def gen_edt() -> Path:
    path = OUT / "03_edt_khepri_software.xlsx"
    wb = Workbook()

    ws = wb.active
    ws.title = "EDT"
    ws["H1"] = APP
    ws["H1"].font = Font(bold=True, size=14)
    for i, (code, name, _) in enumerate(EDT_ITEMS):
        ws.cell(row=3 + i, column=8, value=f"{code} {name}")

    ws2 = wb.create_sheet("Diccionario")
    dheaders = ["Nivel", "Codigo EDT", "Nombre", "Definicion", "Responsable", "Supuestos", "Restricciones", "Riesgos"]
    hf = Font(bold=True, color="FFFFFF")
    fill = PatternFill("solid", fgColor="1565C0")
    for c, h in enumerate(dheaders, 1):
        cell = ws2.cell(1, c, h)
        cell.font = hf
        cell.fill = fill
    for ri, (code, name, defn) in enumerate(EDT_ITEMS, 2):
        ws2.cell(ri, 1, code.count(".") + (0 if code.endswith(".0") else 0))
        ws2.cell(ri, 2, code)
        ws2.cell(ri, 3, name)
        ws2.cell(ri, 4, defn)
        ws2.cell(ri, 5, "Equipo Khepri Software")
        ws2.cell(ri, 6, "Conectividad a internet; credenciales válidas")
        ws2.cell(ri, 7, "Stack FastAPI + React + MySQL")
        ws2.cell(ri, 8, "Scope creep; fallas integración correo")

    ws3 = wb.create_sheet("Recursos")
    ws3.append(["Nombre", "Tipo", "Costo/Hora o Unidad", "Consideraciones"])
    recursos = [
        ("Gerente de proyecto", "Trabajo", 0, "Coordinación"),
        ("Analista / Desarrollador backend", "Trabajo", 0, "FastAPI, MySQL"),
        ("Desarrollador frontend", "Trabajo", 0, "React, TypeScript"),
        ("QA", "Trabajo", 0, "Pruebas UAT"),
        ("Hosting Railway", "Costo", 0, "API + MySQL + frontend"),
        ("Resend (correo)", "Costo", 0, "Plan según volumen"),
    ]
    for r in recursos:
        ws3.append(list(r))

    wb.save(path)
    return path


GANTT_TASKS = [
    ("Kick-off y acta", 0, 5),
    ("Análisis y planilla requerimientos", 5, 15),
    ("Diseño BD y RBAC", 15, 30),
    ("Módulo autenticación", 25, 45),
    ("Módulo empresas y multi-tenant", 40, 55),
    ("Catálogo (productos, bodegas)", 50, 75),
    ("Inventario operativo", 70, 100),
    ("Dashboard y exportaciones", 95, 110),
    ("Pruebas integración", 105, 120),
    ("Despliegue Railway", 115, 125),
    ("Documentación ERS/casos de uso", 120, 130),
    ("Cierre y aceptación", 128, 135),
]


def gen_gantt() -> Path:
    path = OUT / "05_carta_gantt_khepri.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Gantt"
    start = date(2026, 1, 6)
    headers = ["ID", "Tarea", "Inicio", "Fin", "Duración (días)", "Predecesora"]
    hf = Font(bold=True, color="FFFFFF")
    fill = PatternFill("solid", fgColor="1565C0")
    for c, h in enumerate(headers, 1):
        cell = ws.cell(1, c, h)
        cell.font = hf
        cell.fill = fill
    for i, (name, d0, d1) in enumerate(GANTT_TASKS, 1):
        ini = start + timedelta(days=d0)
        fin = start + timedelta(days=d1)
        ws.cell(i + 1, 1, i)
        ws.cell(i + 1, 2, name)
        ws.cell(i + 1, 3, ini.isoformat())
        ws.cell(i + 1, 4, fin.isoformat())
        ws.cell(i + 1, 5, d1 - d0)
        ws.cell(i + 1, 6, i - 1 if i > 1 else "")
    for col, w in zip("ABCDEF", [5, 38, 12, 12, 14, 12]):
        ws.column_dimensions[col].width = w
    wb.save(path)
    return path


TRACE_MAP = {
    "R.01": ("CU001", "Auth", "LoginPage / POST /auth/login", "Implementado"),
    "R.02": ("CU002", "Auth", "ForgotPasswordPage / olvido-contrasena", "Implementado"),
    "R.03": ("CU001", "Auth", "auth_service bloqueo", "Implementado"),
    "R.04": ("CU003", "Empresas", "EmpresasPage", "Implementado"),
    "R.05": ("CU003", "Empresas", "useEmpresaMaestraFilter", "Implementado"),
    "R.06": ("CU003", "Empresas", "provisionar-rbac", "Implementado"),
    "R.18": ("CU006", "Inventario", "InventarioPage stock", "Implementado"),
    "R.29": ("CU007", "Inventario", "POST /inventario/recepcion", "Implementado"),
    "R.36": ("CU010", "Auth", "POST cambiar-contrasena", "Implementado"),
    "R.46": ("CU003", "Empresas", "DELETE /empresas inhabilitar", "Implementado"),
}


def gen_trazabilidad() -> Path:
    path = OUT / "06_matriz_trazabilidad_khepri.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Trazabilidad"
    headers = ["R-N°", "Requerimiento", "Tipo", "Caso de uso", "Módulo", "Artefacto", "Estado impl."]
    hf = Font(bold=True, color="FFFFFF")
    fill = PatternFill("solid", fgColor="1565C0")
    for c, h in enumerate(headers, 1):
        cell = ws.cell(1, c, h)
        cell.font = hf
        cell.fill = fill
    for ri, row in enumerate(ROWS, 2):
        rid = row[0]
        cu, mod, art, est = TRACE_MAP.get(rid, ("—", row[3], "Ver ERS", "Implementado"))
        ws.cell(ri, 1, rid)
        ws.cell(ri, 2, row[1])
        ws.cell(ri, 3, row[2])
        ws.cell(ri, 4, cu)
        ws.cell(ri, 5, mod)
        ws.cell(ri, 6, art)
        ws.cell(ri, 7, est)
    widths = [8, 36, 14, 10, 18, 40, 14]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"
    wb.save(path)
    return path


PANTALLAS = [
    ("P01", "Landing", "/", "Visitante", "R.23"),
    ("P02", "Login", "/login", "Todos", "R.01"),
    ("P03", "Olvido contraseña", "/olvido-contrasena", "Todos", "R.02"),
    ("P04", "Restablecer contraseña", "/restablecer-contrasena", "Todos", "R.02"),
    ("P05", "Dashboard", "/app", "Autenticado", "R.44"),
    ("P06", "Empresas", "/app/empresas", "Super admin", "R.04, R.46"),
    ("P07", "Usuarios", "/app/usuarios", "Admin", "R.07"),
    ("P08", "Productos", "/app/productos", "Admin/Operador", "R.16, R.17"),
    ("P09", "Bodegas", "/app/bodegas", "Admin", "R.11"),
    ("P10", "Inventario", "/app/inventario", "Operador", "R.18–R.21, R.29–R.32"),
    ("P11", "Roles / Permisos", "/app/roles, /app/permisos", "Admin", "R.09, R.38"),
    ("P12", "Perfil", "/app/perfil", "Autenticado", "R.22"),
    ("P13", "404", "/*", "Todos", "R.43"),
]


def gen_catalogo_pantallas() -> Path:
    path = OUT / "07_catalogo_pantallas_khepri.docx"
    doc = Document()
    doc.add_heading(f"Catálogo de pantallas (Mockups as-built) — {APP}", 0)
    _p(doc, "Inventario de interfaces implementadas. Capturas de pantalla pueden adjuntarse en anexo visual.")
    _table(doc, ["ID", "Pantalla", "Ruta", "Actores", "Requerimientos"], PANTALLAS)
    _hdr(doc, "Notas para mockups formales", 2)
    _p(doc, "Para entrega académica: exportar capturas desde localhost:5173 o Railway y pegar en Word/Figma.")
    doc.save(path)
    return path


def gen_minuta() -> Path:
    path = OUT / "08_plantilla_minuta_reunion.docx"
    doc = Document()
    doc.add_heading("Minuta de reunión — Khepri Software", 0)
    _table(doc, ["Campo", "Valor"], [
        ["Fecha", FECHA.strftime("%d/%m/%Y")],
        ["Lugar / Modalidad", "Presencial / Remoto"],
        ["Convocante", ""],
        ["Objetivo de la reunión", ""],
    ])
    _hdr(doc, "Asistentes", 2)
    _table(doc, ["Nombre", "Cargo", "Firma"], [["", "", ""]] * 4)
    _hdr(doc, "Temas tratados", 2)
    _table(doc, ["#", "Tema", "Detalle", "Responsable"], [["1", "", "", ""]] * 3)
    _hdr(doc, "Acuerdos y compromisos", 2)
    _table(doc, ["#", "Acuerdo", "Fecha compromiso", "Responsable"], [["1", "", "", ""]] * 3)
    _hdr(doc, "Próxima reunión", 2)
    _p(doc, "Fecha: __________  Hora: __________")
    doc.save(path)
    return path


def gen_plan_pruebas() -> Path:
    path = OUT / "09_plan_pruebas_uat_khepri.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "UAT"
    headers = ["ID", "Caso de prueba", "Pasos", "Resultado esperado", "R-N°", "Estado"]
    hf = Font(bold=True, color="FFFFFF")
    fill = PatternFill("solid", fgColor="1565C0")
    for c, h in enumerate(headers, 1):
        cell = ws.cell(1, c, h)
        cell.font = hf
        cell.fill = fill
    casos = [
        ("TP01", "Login válido", "Ingresar credenciales correctas", "Acceso al panel /app", "R.01", "Pendiente"),
        ("TP02", "Login inválido", "3 intentos fallidos", "Bloqueo temporal", "R.03", "Pendiente"),
        ("TP03", "Cambio contraseña perfil", "Usuario cambia contraseña en /app/perfil", "Contraseña actualizada", "R.02", "Pendiente"),
        ("TP04", "Recepción stock", "Recepcionar 10 unidades SKU X", "Stock incrementado", "R.29", "Pendiente"),
        ("TP05", "Inhabilitar empresa", "Super admin inhabilita tenant", "Oculta en agregados", "R.46", "Pendiente"),
        ("TP06", "Exportar stock Excel", "Descargar desde inventario", "Archivo xlsx válido", "R.47", "Pendiente"),
        ("TP07", "Carga masiva productos", "Importar plantilla Excel", "Productos creados", "R.17", "Pendiente"),
        ("TP08", "Permiso denegado", "Usuario sin permiso accede módulo", "403 o UI bloqueada", "R.25", "Pendiente"),
    ]
    for ri, row in enumerate(casos, 2):
        for ci, val in enumerate(row, 1):
            ws.cell(ri, ci, val)
    wb.save(path)
    return path


def gen_indice() -> Path:
    path = OUT / "00_indice_documentacion.md"
    content = f"""# Documentación de proyecto — {APP}

Generado: {FECHA.isoformat()}

| # | Documento | Archivo |
|---|-----------|---------|
| 0 | Índice | `00_indice_documentacion.md` |
| 1 | Acta de constitución | `01_acta_constitucion_khepri.docx` |
| 2 | ERS | `02_ers_khepri_software.docx` |
| 3 | EDT (WBS + diccionario + recursos) | `03_edt_khepri_software.xlsx` |
| 4 | Casos de uso (alto nivel + extendidos) | `04_casos_uso_khepri.docx` |
| 5 | Carta Gantt | `05_carta_gantt_khepri.xlsx` |
| 6 | Matriz de trazabilidad | `06_matriz_trazabilidad_khepri.xlsx` |
| 7 | Catálogo de pantallas / mockups | `07_catalogo_pantallas_khepri.docx` |
| 8 | Plantilla minuta de reunión | `08_plantilla_minuta_reunion.docx` |
| 9 | Plan de pruebas UAT | `09_plan_pruebas_uat_khepri.xlsx` |
| 10 | Planilla 50 requerimientos | `10_planilla_requerimientos_khepri.xlsx` |

Regenerar todo:

```bash
python scripts/generar_documentacion_proyecto.py
python scripts/generar_requerimientos_excel.py
```
"""
    path.write_text(content, encoding="utf-8")
    return path


def main() -> None:
    gen_requerimientos()
    req_src = ROOT / "docs" / "requerimientos_khepri_software.xlsx"
    req_dst = OUT / "10_planilla_requerimientos_khepri.xlsx"
    shutil.copy2(req_src, req_dst)
    files = [
        gen_indice(),
        gen_acta(),
        gen_ers(),
        gen_edt(),
        gen_casos_uso(),
        gen_gantt(),
        gen_trazabilidad(),
        gen_catalogo_pantallas(),
        gen_minuta(),
        gen_plan_pruebas(),
    ]
    print(f"Documentación generada en: {OUT}\n")
    for f in files:
        print(f"  - {f.name}")
    print(f"  - 10_planilla_requerimientos_khepri.xlsx (copia en proyecto/)")


if __name__ == "__main__":
    main()
